
from docx import Document
import textblob
import nltk
from collections import Counter



f = open("KT_New_v1.docx", "rb")
doc = Document(f)


paragraphs = [x.text for x in doc.paragraphs if x != '']

all = "\n".join([x.replace("\t","") for x in paragraphs])

blob = textblob.TextBlob(all)

#stemmer = nltk.PorterStemmer()
stemmer = nltk.stem.WordNetLemmatizer()

# Sections (doc.sections): 1
print("Sections:", len(doc.sections))

# Paragraphs (doc.paragraphs): 295
print("Paragraphs:", len(doc.paragraphs))

# Sentences (blob): 35
print("Sentences:", len(blob.sentences))

# Words: (blob.words): 1617
print("Words:", len(blob.words))

# Tokens: (blob.tokenize()): 2113
print("Tokens:", len(blob.tokens))

# Verbs
verbs = [k.lower() for k,v in blob.tags if v.startswith("V")]
print("Verbs:", len(verbs))

lemmatized = []
for v in verbs:
    lemmatized.append(stemmer.lemmatize(v, 'v'))


print(Counter(lemmatized))
# Subjectivity
print(f'Subjectivity: {round(blob.subjectivity*100,2)}%')
print(f'Objectivity: {100-round(blob.subjectivity*100,2)}%')

# Sentiment
print(blob.sentiment)
